from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path('/Users/epower/Documents/New project/共享单车大数据分析与可视化技术调研报告_结构化合规版.docx')

NAVY = '17365D'
BLUE = '1F4E79'
LIGHT_BLUE = 'DCEAF7'
PALE_BLUE = 'F4F8FC'
GRID = 'D9D9D9'
GRAY = '595959'
BLACK = '000000'


def set_run_font(run, size=None, bold=None, color=None, font='Noto Sans SC'):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    rfonts.set(qn('w:eastAsia'), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, color=GRID, size='6'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run('第 ')
    set_run_font(run, 9, color=GRAY)
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(' 页')
    set_run_font(tail, 9, color=GRAY)


def set_para_format(p, before=0, after=7, line=1.55, first=0.74):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first:
        pf.first_line_indent = Cm(first)


def add_body(doc, text, before=0, after=7):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_format(p, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, 11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.55)
    p.paragraph_format.first_line_indent = Cm(-0.37)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    set_run_font(r, 10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    p.paragraph_format.space_before = Pt(15 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, 15 if level == 1 else 12, bold=True, color=BLACK, font='Noto Sans SC')
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, label in enumerate(headers):
        cell = header.cells[i]
        cell.text = ''
        set_cell_shading(cell, NAVY)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, 10, bold=True, color='FFFFFF', font='Noto Sans SC')
        if widths:
            cell.width = Cm(widths[i])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cell = cells[i]
            cell.text = ''
            if row_index % 2 == 1:
                set_cell_shading(cell, PALE_BLUE)
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(value) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            r = p.add_run(value)
            set_run_font(r, 9.5, color=BLACK)
            if widths:
                cell.width = Cm(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def configure_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Noto Sans SC'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans SC')
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.55
    normal.paragraph_format.space_after = Pt(7)

    for name, size in [('Title', 24), ('Subtitle', 13), ('Heading 1', 15), ('Heading 2', 12)]:
        style = doc.styles[name]
        style.font.name = 'Noto Sans SC'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = name != 'Subtitle'


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(102)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('共享单车大数据分析与可视化技术应用调研报告')
    set_run_font(r, 25, bold=True, color=BLACK, font='Noto Sans SC')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(52)
    r = p.add_run('面向课程项目的技术路线  应用场景与合规性核查')
    set_run_font(r, 13, color=GRAY)

    info = [
        ('课程名称', '软件开发实践1-25-26-2'),
        ('指导教师', '葛瑞泉'),
        ('项目主题', '共享单车租赁需求分析  数据库查询与可视化大屏系统'),
        ('项目小组', '邹晓谦  赵轩逸  侯季南  高轩  王浩轩'),
        ('提交日期', '2026年9月'),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, (label, value) in enumerate(info):
        for j, text in enumerate((label, value)):
            cell = table.cell(i, j)
            cell.width = Cm(3.4 if j == 0 else 10.2)
            set_cell_border(cell, color='FFFFFF', size='0')
            set_cell_margins(cell, top=95, start=90, bottom=95, end=90)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, 11, bold=(j == 0), color=GRAY if j == 0 else BLACK, font='Noto Sans SC')
    doc.add_page_break()


def add_header_footer(doc):
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.top_margin = Cm(2.1)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run('共享单车大数据分析与可视化技术应用调研报告')
        set_run_font(r, 8.5, color=GRAY)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(p)
        # Keep the cover page free of running headers and page numbers.
        section.first_page_header.paragraphs[0].text = ''
        section.first_page_footer.paragraphs[0].text = ''


def build():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)

    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('摘要')
    set_run_font(r, 15, bold=True, font='Noto Sans SC')
    add_body(doc, '本报告以共享单车租赁需求分析、数据库查询与可视化大屏系统为对象，围绕公开数据使用、数据处理、数据库查询、可视化展示及质量验证开展技术调研。项目采用 UCI Machine Learning Repository 发布的 Bike Sharing Dataset 作为小时级统计分析的主要数据来源，并将本地整理的 Kaggle Bike Sharing Demand 数据用于 SQLite 查询与前端演示。技术实现覆盖 pandas 与 NumPy 数据处理、SQLite 数据存储与查询接口、matplotlib 静态图表、ECharts 前端大屏以及 pytest 自动化测试。')
    add_body(doc, '调研结果表明，小时粒度数据能够识别通勤高峰和天气条件下的租赁量差异；将分析模块、查询模块和展示模块分层，有利于课程项目的复现、测试和答辩说明。合规核查显示：当前课程项目仅处理公开、聚合的示例数据，不包含可直接识别个人身份的真实轨迹数据；但若未来接入真实订单、定位或用户数据，必须遵守个人信息保护、数据安全、数据授权和开源许可证等要求。', after=9)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run('关键词：')
    set_run_font(r, 10.5, bold=True, font='Noto Sans SC')
    r = p.add_run('共享单车；数据分析；SQLite；数据可视化；个人信息保护')
    set_run_font(r, 10.5)

    add_heading(doc, '一、调研背景与目标')
    add_heading(doc, '1.1 调研背景', 2)
    add_body(doc, '共享单车是依托互联网平台提供分时租赁服务的城市慢行交通方式。其运行过程会形成时间、天气、使用频次、用户类型和车辆状态等多维数据。对这类数据进行规范处理，可以帮助研究者理解租赁需求的时间规律，并为运维调度、慢行交通规划和用户服务改进提供描述性依据。')
    add_body(doc, '本调研不以单一平台的商业指标为结论依据，也不对企业运营能力作未经证实的评价；行业应用部分聚焦公开资料能够支持的技术能力与治理问题。')
    add_heading(doc, '1.2 调研目标与边界', 2)
    add_bullet(doc, '梳理共享单车数据从采集、清洗、存储、查询到可视化展示的实现链路。')
    add_bullet(doc, '说明本课程项目所采用的数据来源、字段范围、分析方法与测试方式。')
    add_bullet(doc, '以项目统计结果为基础，区分描述性发现、合理解释和后续扩展方向，避免将相关性表述为因果关系。')
    add_bullet(doc, '核查数据来源、个人信息处理、平台素材与开源组件在课程项目场景下的主要合规要求。')

    add_heading(doc, '二、数据来源与研究方法')
    add_heading(doc, '2.1 数据来源及使用范围', 2)
    add_table(doc,
              ['数据集', '记录规模与字段', '在项目中的用途', '授权或使用注意事项'],
              [
                  ['UCI Bike Sharing Dataset', 'hour.csv 本地文件 17,379 条；含小时、季节、工作日、天气、温湿度、临时用户、注册用户和总租车量。', '用于清洗、分组统计和静态图表。', 'UCI 页面标注为 CC BY 4.0；提交报告与代码时保留来源及署名。'],
                  ['Kaggle Bike Sharing Demand', 'bike_data.csv 本地文件 10,886 条；含时间、季节、天气、温湿度和租车量等字段。', '用于 SQLite 建库、查询接口和前端 JSON 演示。', '仅按课程演示范围使用；对外发布前应以原数据页面的许可和竞赛规则为准。'],
              ], widths=[3.2, 4.4, 3.6, 4.2])
    add_heading(doc, '2.2 数据处理与统计方法', 2)
    add_body(doc, '数据处理遵循“读取与字段校验—缺失值和异常值检查—衍生时间字段—分组聚合—结果复核”的流程。针对 UCI 小时数据，以 hr、workingday、weathersit、casual、registered 和 cnt 等字段作为核心分析变量；针对 Kaggle 数据，以 datetime、weather、temp、casual、registered 和 count 等字段支撑数据库查询与大屏指标。')
    add_body(doc, '统计方法以描述性分析为主：按小时计算平均租车量，按天气编码比较租赁量均值，并比较工作日与非工作日、注册用户与临时用户的构成差异。结果仅反映样本内分布，不用于推断因果关系或直接外推至其他城市。')
    add_heading(doc, '2.3 项目技术路线', 2)
    add_table(doc,
              ['层级', '主要技术', '职责与交付'],
              [
                  ['分析层', 'Python、pandas、NumPy、matplotlib', '完成 CSV 读取、字段校验、清洗、统计指标与 PNG 图表输出。'],
                  ['数据服务层', 'SQLite、Python 查询模块', '创建 bike_usage 表及索引，提供 KPI、天气和小时维度查询结果。'],
                  ['展示层', 'HTML、CSS、JavaScript、ECharts', '将处理后的 data.json 渲染为指标卡片和趋势图表。'],
                  ['质量保障层', 'pytest', '验证数据读取、清洗后记录数、统计指标与数据库导入流程。'],
              ], widths=[2.2, 4.4, 8.8])

    add_heading(doc, '三、关键技术与系统实现')
    add_heading(doc, '3.1 数据清洗与统计分析', 2)
    add_body(doc, 'pandas 提供适用于结构化表格数据的 DataFrame 结构，可完成 CSV 读取、类型转换、分组聚合和结果导出；NumPy 为数值计算提供基础支持。项目将数据清洗与指标计算拆分为独立函数，以减少前端展示逻辑与分析逻辑的耦合。')
    add_heading(doc, '3.2 SQLite 建库与查询接口', 2)
    add_body(doc, '项目将 Kaggle 整理数据导入 SQLite 的 bike_usage 表，并围绕答辩展示设计三类查询：总租车量、注册用户量和平均温度等 KPI 汇总；按天气编码聚合租车量；按小时计算平均租车量。查询结果以 Python 字典或 JSON 结构返回，便于前端复用。')
    add_heading(doc, '3.3 可视化与测试', 2)
    add_body(doc, 'matplotlib 负责生成适合报告引用的静态折线图和柱状图；ECharts 负责浏览器端数据大屏。pytest 用于核对输入数据、清洗结果和关键指标，使从原始 CSV 到图表或接口结果的过程可重复执行。')

    add_heading(doc, '四、分析结果与专业解读')
    add_heading(doc, '4.1 主要统计发现', 2)
    add_table(doc,
              ['分析维度', '项目结果', '解读边界'],
              [
                  ['小时需求', '17 时平均租车量约 461.45 次，为样本内最高时段。', '反映本数据集中的通勤高峰特征，不代表所有城市或所有季节。'],
                  ['天气差异', '天气编码从 1 到 4 时，平均租车量约为 204.87、175.17、111.58、74.33 次。', '天气越差的分组租车量越低，属于描述性关联，未控制节假日、季节等变量。'],
                  ['用户结构', '注册用户租赁量占比约 81.17%。', '表明样本中注册用户贡献更高；不能据此推断用户偏好或用户价值。'],
                  ['工作日差异', '工作日平均租车量约 193.21 次，非工作日约 181.41 次。', '差异需要结合城市、天气与时间段进一步检验。'],
              ], widths=[2.5, 6.4, 6.5])
    add_heading(doc, '4.2 对系统设计的启示', 2)
    add_body(doc, '分析结果说明，面向运营展示的指标至少应同时提供时间维度、天气维度和用户结构维度。系统将“静态分析—数据库查询—前端展示”分层实现，使分析脚本可独立复跑，查询接口可独立测试，展示页面可直接消费 JSON 数据。该结构更适合作为后续需求预测、异常检测和实时数据接入的基础。')

    add_heading(doc, '五、应用场景与技术发展方向')
    add_heading(doc, '5.1 可落地的应用场景', 2)
    add_bullet(doc, '运维调度：以小时、天气和区域等维度识别可能的供需不均衡时段，为人工调度提供参考。')
    add_bullet(doc, '慢行交通评估：将汇总数据与公共交通站点、道路设施等合法取得的外部数据结合，评估接驳需求。')
    add_bullet(doc, '绿色出行研究：在明确排放核算口径和替代出行假设的前提下，估算骑行活动的环境效益。')
    add_heading(doc, '5.2 技术演进方向', 2)
    add_body(doc, '后续可将当前描述性分析扩展为需求预测模块，并采用 MAE、RMSE 和 R² 等指标评价模型。对于实时运营场景，可在数据授权和安全措施完备的前提下，引入流数据处理、时空索引和异常检测；对于多源融合场景，应先统一数据口径、时间粒度和权限边界。')

    add_heading(doc, '六、数据与合法合规性核查')
    add_heading(doc, '6.1 核查结论', 2)
    add_body(doc, '本项目在当前课程演示范围内具备可控的合规基础：分析主数据来自公开 UCI 数据集，数据字段为租赁量、天气和时间等聚合变量；项目未接入真实用户身份、手机号、精确位置轨迹或支付信息。UCI 数据集页面标明 CC BY 4.0 许可，允许在保留适当署名的条件下共享和改编。')
    add_body(doc, '但该结论仅适用于本报告列明的公开样例数据和课程项目用途。若改用真实运营数据、接入第三方接口、公开部署系统或用于商业决策，需重新开展数据来源、处理目的、权限、留存周期和安全措施的专项审查。')
    add_heading(doc, '6.2 个人信息与数据安全要求', 2)
    add_body(doc, '《中华人民共和国个人信息保护法》将行踪轨迹列为敏感个人信息。真实共享单车订单中的精确起止位置、时间序列、设备标识和账户信息在组合后可能识别或推断个人出行行为。因此，未来处理真实数据时，应先确定合法处理基础，遵循目的明确、最小必要、公开透明和安全保障原则；处理敏感个人信息通常还需要满足特定目的、充分必要性、严格保护措施及单独同意等要求。')
    add_body(doc, '对面向真实用户的需求预测、动态定价或信用限制功能，还应评估自动化决策的透明度与公平性。涉及敏感个人信息、自动化决策、对外提供或跨境提供个人信息等情形时，应按法律要求开展个人信息保护影响评估并留存记录。')
    add_heading(doc, '6.3 数据来源 版权与开源组件', 2)
    add_table(doc,
              ['核查项目', '核查结果', '提交与后续使用要求'],
              [
                  ['UCI 数据集', '页面明确标注 CC BY 4.0。', '在报告、代码仓库和数据说明中保留数据集名称、作者、链接和许可信息。'],
                  ['Kaggle 整理数据', '项目中已用于本地演示；本报告未将其许可视为当然授权。', '在任何对外发布或再次分发前，核对原数据页的许可、竞赛规则和来源声明。'],
                  ['ECharts', '官方仓库以 Apache License 2.0 发布。', '分发前端代码时保留许可证、版权和 NOTICE 要求；不暗示与 Apache 的官方关联。'],
                  ['截图与第三方素材', '数据页面截图仅用于课程过程证据。', '标明来源；不将平台页面、商标或非授权图片作为可再分发素材。'],
              ], widths=[2.8, 5.2, 7.4])
    add_heading(doc, '6.4 合规清单', 2)
    add_bullet(doc, '课程提交前：核对数据说明、许可证和参考文献链接，删除未核实的行业规模、企业内部指标和来源不明图片。')
    add_bullet(doc, '代码交付时：不提交真实用户标识、密钥、账号、定位轨迹或含敏感信息的日志；保留依赖与开源许可说明。')
    add_bullet(doc, '系统扩展前：建立数据分级、访问控制、脱敏或聚合策略、留存和删除规则，并按具体业务重新评估法律要求。')

    add_heading(doc, '七、结论与建议')
    add_body(doc, '本调研围绕共享单车数据分析与可视化项目，形成了从公开数据使用、清洗统计、数据库查询、可视化展示到自动化测试的完整技术说明。项目结果能够支持对小时高峰、天气差异和用户结构的描述性分析，并通过分层架构提高了流程的可复现性与可解释性。')
    add_body(doc, '建议后续工作保持三项原则：第一，数据结论与数据边界同时呈现，避免过度外推；第二，所有数据、代码和素材均保留可追溯来源与许可信息；第三，在接入真实运营数据或部署面向用户的功能前，先完成个人信息、数据安全和自动化决策方面的专项评估。')

    add_heading(doc, '参考文献')
    references = [
        '［1］Fanaee-T, H. Bike Sharing Dataset［DB/OL］. UCI Machine Learning Repository, 2013. https://doi.org/10.24432/C5W894.',
        '［2］中华人民共和国个人信息保护法［Z/OL］. 中国人大网, 2021. https://www.npc.gov.cn/npc/c2/c30834/202108/t20210820_3130888.html.',
        '［3］交通运输部等十部门. 关于鼓励和规范互联网租赁自行车发展的指导意见［Z/OL］. 交通运输部政府信息公开平台. https://xxgk.mot.gov.cn/jigouapp/741/788/838/list_6699.html.',
        '［4］Apache ECharts. License［EB/OL］. https://github.com/apache/echarts/blob/master/LICENSE.',
        '［5］pandas development team. pandas documentation［EB/OL］. https://pandas.pydata.org/docs/.',
        '［6］pytest contributors. pytest documentation［EB/OL］. https://docs.pytest.org/.',
    ]
    for ref in references:
        p = doc.add_paragraph()
        set_para_format(p, after=4, line=1.25, first=0)
        r = p.add_run(ref)
        set_run_font(r, 9.5)

    # Set core properties without retaining source-author metadata.
    props = doc.core_properties
    props.title = '共享单车大数据分析与可视化技术应用调研报告'
    props.subject = '课程项目技术调研与合规性核查'
    props.author = '课程项目小组'
    props.keywords = '共享单车, 数据分析, 可视化, 合规'
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
